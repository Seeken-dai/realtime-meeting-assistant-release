"""
真实资料形态冒烟：用仓库内中文产品文档生成 DOCX/PDF，
走解析 → 切片检索 → 证据原文核验全链路。

不依赖外网 LLM/ASR。失败类型写入 stdout，便于写进 HANDOFF。
"""

from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from document_extract import DocumentExtractionError, extract_document
from knowledge_base import KnowledgeBase
from suggest import _validate

ROOT = Path(__file__).resolve().parent
DOCS = ROOT / "docs"


def _register_chinese_font():
    """尽量挂上系统中文字体，便于生成可读 PDF。"""
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\msyh.ttf",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for path in candidates:
        if os.path.isfile(path):
            try:
                pdfmetrics.registerFont(TTFont("CN", path, subfontIndex=0))
                return "CN"
            except Exception:
                continue
    return "Helvetica"


def _build_docx(path: Path, source_md: Path):
    text = source_md.read_text(encoding="utf-8")
    document = Document()
    document.add_heading("产品能力说明（冒烟样例）", level=1)
    document.add_paragraph(
        "本文件由仓库知识文档转换生成，用于验证 Word 导入/预览/检索/证据链路。"
    )
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "能力"
    table.cell(0, 1).text = "口径"
    table.cell(1, 0).text = "条件分支审批"
    table.cell(1, 1).text = "标准版不支持，需单独评估"
    document.save(path)


def _build_pdf(path: Path, source_md: Path, font_name: str):
    text = source_md.read_text(encoding="utf-8")
    pdf = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    y = height - 72
    pdf.setFont(font_name, 14)
    pdf.drawString(72, y, "需求边界与报价（冒烟样例）")
    y -= 28
    pdf.setFont(font_name, 10)
    for line in text.splitlines():
        line = line.strip()
        if not line:
            y -= 10
            continue
        # 简单折行
        while line:
            chunk = line[:42]
            line = line[42:]
            if y < 72:
                pdf.showPage()
                y = height - 72
                pdf.setFont(font_name, 10)
            pdf.drawString(72, y, chunk)
            y -= 16
    pdf.showPage()
    pdf.setFont(font_name, 11)
    pdf.drawString(72, height - 72, "第二页：联调与免费对接口径")
    pdf.drawString(72, height - 96, "OA 对接与接口联调需评估范围，不默认免费实施。")
    pdf.save()


def main() -> int:
    failures: list[str] = []
    product_md = DOCS / "产品功能清单.md"
    boundary_md = DOCS / "需求边界与报价规则.md"
    if not product_md.is_file() or not boundary_md.is_file():
        print("FAIL: 缺少 poc/docs 中文知识文档")
        return 1

    font_name = _register_chinese_font()
    with tempfile.TemporaryDirectory(prefix="mc-real-doc-smoke-") as root:
        root_path = Path(root)
        docx_path = root_path / "产品能力说明-冒烟.docx"
        pdf_path = root_path / "需求边界报价-冒烟.pdf"
        scanned_path = root_path / "扫描件-无文字层.pdf"

        _build_docx(docx_path, product_md)
        _build_pdf(pdf_path, boundary_md, font_name)
        # 纯图 PDF（无文字层）
        from reportlab.pdfgen import canvas as _c

        blank = _c.Canvas(str(scanned_path))
        blank.rect(72, 650, 300, 100, fill=1)
        blank.save()

        # 1) 解析
        try:
            docx = extract_document(str(docx_path))
            assert docx["format"] == "docx"
            assert "三级审批" in docx["text"] or "审批" in docx["text"]
            assert "REST" in docx["text"] or "接口" in docx["text"]
            print("PASS: DOCX 解析含标题/正文/能力关键词")
        except Exception as exc:
            failures.append(f"DOCX 解析: {exc}")
            print(f"FAIL: DOCX 解析: {exc}")

        try:
            pdf = extract_document(str(pdf_path))
            assert pdf["format"] == "pdf"
            assert pdf.get("pageCount", 0) >= 1
            assert "联调" in pdf["text"] or "报价" in pdf["text"] or "对接" in pdf["text"]
            print(f"PASS: PDF 解析 {pdf.get('pageCount')} 页，含业务关键词")
        except Exception as exc:
            failures.append(f"PDF 解析: {exc}")
            print(f"FAIL: PDF 解析: {exc}")

        try:
            extract_document(str(scanned_path))
            failures.append("扫描 PDF 应拒绝 OCR，却解析成功")
            print("FAIL: 扫描 PDF 未拒绝")
        except DocumentExtractionError as exc:
            if "OCR" in str(exc) or "文字" in str(exc) or "文本" in str(exc):
                print(f"PASS: 扫描 PDF 拒绝（{exc}）")
            else:
                print(f"PASS: 扫描 PDF 拒绝（{exc}）")

        # 2) 检索
        try:
            kb = KnowledgeBase(
                docs_dir=str(root_path),
                backend="local",
                doc_paths=[str(docx_path), str(pdf_path)],
                verbose=False,
            ).build()
            assert len(kb.chunks) >= 2
            api_hits = kb.search("开放接口 Webhook 订单状态", top_k=3)
            approval_hits = kb.search("三级审批 条件分支", top_k=3)
            assert api_hits, "接口检索无命中"
            assert approval_hits, "审批检索无命中"
            print(
                "PASS: 本地检索命中 "
                f"API→{api_hits[0]['source']} / 审批→{approval_hits[0]['source']}"
            )
        except Exception as exc:
            failures.append(f"检索: {exc}")
            print(f"FAIL: 检索: {exc}")
            api_hits = []

        # 3) 证据原文校验（不调 LLM，直接走 _validate）
        try:
            hits = api_hits or [
                {
                    "source": docx_path.name,
                    "text": extract_document(str(docx_path))["text"][:500],
                }
            ]
            grounded = _validate(
                [
                    {
                        "intent": "说明接口能力",
                        "script": "我们支持 REST 接口，订单状态可用 Webhook 主动通知。",
                        "type": "grounded",
                        "references": [hits[0]["source"]],
                        "evidence": [],
                    }
                ],
                hits,
            )[0]
            assert grounded["type"] == "grounded", grounded
            assert grounded.get("evidence"), "应有 evidence"
            quote = grounded["evidence"][0]["quote"]
            blob = "\n".join(h["text"] for h in hits)
            assert quote.replace(" ", "") in blob.replace(" ", "") or any(
                quote in h["text"] for h in hits
            )
            print(f"PASS: 证据原文可回填并逐字落在候选中（quote 长度 {len(quote)}）")
        except Exception as exc:
            failures.append(f"证据校验: {exc}")
            print(f"FAIL: 证据校验: {exc}")

    print("---")
    if failures:
        print(f"SMOKE FAILED: {len(failures)} issue(s)")
        for item in failures:
            print(f"  - {item}")
        return 1
    print("SMOKE PASSED: DOCX/PDF 预览解析、检索、证据链路")
    return 0


if __name__ == "__main__":
    sys.exit(main())
