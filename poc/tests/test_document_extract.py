"""DOCX/PDF 知识解析、切片与检索的无网络回归测试。"""

import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from reportlab.pdfgen import canvas

from document_extract import DocumentExtractionError, extract_document
from knowledge_base import KnowledgeBase


def _make_docx(path):
    document = Document()
    document.add_heading("Product boundaries", level=1)
    document.add_paragraph("Standard edition supports a fixed three-level approval flow.")
    document.add_heading("Delivery", level=2)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Policy"
    table.cell(1, 0).text = "Custom workflow"
    table.cell(1, 1).text = "Requires separate assessment"
    document.save(path)


def _make_pdf(path, with_text=True):
    pdf = canvas.Canvas(path)
    if with_text:
        pdf.setFont("Helvetica-Bold", 16)
        pdf.drawString(72, 760, "PDF Product Guide")
        pdf.setFont("Helvetica", 11)
        pdf.drawString(72, 730, "Standard edition includes twelve reports.")
        pdf.showPage()
        pdf.drawString(72, 760, "Mobile app supports approvals and notifications.")
    else:
        pdf.rect(72, 650, 300, 100, fill=1)
    pdf.save()


def test_docx_pdf_extraction_and_search():
    with tempfile.TemporaryDirectory(prefix="meeting-doc-extract-") as root:
        docx_path = os.path.join(root, "product.docx")
        pdf_path = os.path.join(root, "guide.pdf")
        scanned_path = os.path.join(root, "scan.pdf")
        _make_docx(docx_path)
        _make_pdf(pdf_path)
        _make_pdf(scanned_path, with_text=False)

        docx = extract_document(docx_path)
        assert docx["format"] == "docx"
        assert "# Product boundaries" in docx["text"]
        assert "fixed three-level approval flow" in docx["text"]
        assert "| Item | Policy |" in docx["text"]

        pdf = extract_document(pdf_path)
        assert pdf["format"] == "pdf"
        assert pdf["pageCount"] == 2
        assert "## 第 1 页" in pdf["text"]
        assert "twelve reports" in pdf["text"]
        assert "Mobile app supports approvals" in pdf["text"]

        try:
            extract_document(scanned_path)
            raise AssertionError("scanned PDF should require OCR")
        except DocumentExtractionError as exc:
            assert "OCR" in str(exc)

        kb = KnowledgeBase(
            docs_dir=root,
            backend="local",
            doc_paths=[docx_path, pdf_path],
            verbose=False,
        ).build()
        assert len(kb.chunks) >= 4
        docx_hits = kb.search("three-level approval flow", top_k=2)
        assert docx_hits and docx_hits[0]["source"] == "product.docx"
        pdf_hits = kb.search("twelve reports", top_k=2)
        assert pdf_hits and pdf_hits[0]["source"] == "guide.pdf"


if __name__ == "__main__":
    test_docx_pdf_extraction_and_search()
    print("ok: docx/pdf extraction + scanned PDF rejection + local retrieval")
